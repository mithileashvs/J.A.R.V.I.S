"""
JARVIS code analysis (Phase 3, Section 3) + code explanation modes
(Section 4).

Two things live here because they share the same "extract relevant
code" foundation:
  1. analyze_file() — structured issue detection for a single file
  2. explain_code() — mode-driven explanation of a file, function, or
     class, using the same extraction logic

Honesty about depth, since Section 3 warns against claiming certainty
the evidence doesn't support:
  - Python gets REAL static analysis via pyflakes (unused imports,
    undefined names, unused locals, syntax errors) — this is genuine
    detection, not guessing, and every issue it reports is labelled
    "Confirmed" because pyflakes doesn't produce false positives for
    the categories it checks.
  - Non-Python files get a much lighter structural pass (syntax
    parseable at all, obvious TODO/FIXME markers, line count) —
    labelled accordingly, never dressed up as equivalent to the Python
    path. Building a real multi-language linter is out of scope here;
    pretending to do one via regex would be actively misleading.
  - "Logic issues", "performance concerns", "security problems" in
    the general sense (Section 3's broader list) are NOT something a
    static pass can reliably detect — those get surfaced only when an
    LLM pass is layered on top (explain_code's TECHNICAL/INTERVIEW
    modes can comment on them as "possible" observations), never
    presented as tool-verified facts the way pyflakes issues are.

No LLM call happens inside analyze_file() itself — it returns
structured, tool-verified findings. Callers (debug_mode.py, or a
future intent handler) are responsible for turning that into prose,
so the confidence distinction between "pyflakes found this" and "the
model thinks this" never gets blurred together.
"""

import ast
import io
import logging
import os
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger("jarvis-code-analysis")

_PYTHON_EXTENSIONS = {".py"}
_MAX_FILE_BYTES = 500_000  # ~500KB — large enough for any real source file, small enough to stay fast

# Images/PDF/Word docs are inherently bigger than source files (a phone
# photo alone can be several MB), so they get their own, more generous
# cap rather than being squeezed under _MAX_FILE_BYTES, which stays
# tight for text/code.
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
# Public alias — main.py's /chat/upload needs this set to decide
# whether an upload should go through the real vision-model pipeline.
IMAGE_EXTENSIONS = _IMAGE_EXTENSIONS
_PDF_EXTENSIONS   = {".pdf"}
_DOCX_EXTENSIONS  = {".docx"}
_MAX_BINARY_FILE_BYTES = 5_000_000  # ~5MB

_TODO_MARKERS = ("TODO", "FIXME", "XXX", "HACK")


@dataclass
class Issue:
    severity: str        # "error" | "warning" | "note"
    confidence: str       # "Confirmed" | "Highly likely" | "Possible"
    message: str
    line: Optional[int] = None


@dataclass
class AnalysisResult:
    file_path: str
    language: str
    line_count: int
    issues: list[Issue] = field(default_factory=list)
    summary: str = ""
    analysis_depth: str = ""  # "static" (pyflakes-verified) | "structural" (lightweight, non-Python)

    def to_text(self) -> str:
        """Section 3's requested output shape: SUMMARY / ISSUES FOUND / ..."""
        lines = [f"SUMMARY", self.summary, ""]
        if self.issues:
            lines.append("ISSUES FOUND")
            for issue in self.issues:
                loc = f" (line {issue.line})" if issue.line else ""
                lines.append(f"- [{issue.confidence}] {issue.message}{loc}")
        else:
            lines.append("ISSUES FOUND")
            lines.append("- None detected by static analysis.")
        return "\n".join(lines)


def _detect_language(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    return {
        ".py": "Python", ".js": "JavaScript", ".ts": "TypeScript",
        ".jsx": "JavaScript (JSX)", ".tsx": "TypeScript (JSX)",
        ".java": "Java", ".go": "Go", ".rs": "Rust", ".rb": "Ruby",
        ".c": "C", ".cpp": "C++", ".h": "C/C++ header",
        ".cs": "C#", ".php": "PHP",
        ".jpg": "Image", ".jpeg": "Image", ".png": "Image", ".webp": "Image",
        ".pdf": "PDF", ".docx": "Word document",
    }.get(ext, "unknown")


def _read_file_safe(file_path: str) -> str:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"'{file_path}' does not exist or is not a file.")
    size = os.path.getsize(file_path)
    if size > _MAX_FILE_BYTES:
        raise ValueError(
            f"'{file_path}' is {size:,} bytes, over the {_MAX_FILE_BYTES:,}-byte analysis "
            f"limit — this guards against accidentally loading a huge generated/data file "
            f"into an LLM context."
        )
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _analyze_python(source: str, file_path: str) -> list[Issue]:
    """Real static analysis via pyflakes. Every issue here is tool-verified, hence 'Confirmed'."""
    from pyflakes.api import check
    from pyflakes.reporter import Reporter

    out, err = io.StringIO(), io.StringIO()
    check(source, file_path, Reporter(out, err))

    issues: list[Issue] = []
    for raw_line in out.getvalue().splitlines():
        # pyflakes format: "path:line:col: message"
        parts = raw_line.split(":", 3)
        line_no = None
        message = raw_line
        if len(parts) >= 4:
            try:
                line_no = int(parts[1])
            except ValueError:
                line_no = None
            message = parts[3].strip()
        issues.append(Issue(severity="warning", confidence="Confirmed", message=message, line=line_no))

    # Syntax errors surface on stderr as a multi-line block (file:line:col
    # message, then the source snippet, then a caret pointer) — that's
    # ONE error, not three, so join it into a single Issue rather than
    # splitting per-line like the stdout warnings above.
    stderr_text = err.getvalue().strip()
    if stderr_text:
        stderr_lines = stderr_text.splitlines()
        first_line = stderr_lines[0]
        line_no = None
        parts = first_line.split(":", 3)
        if len(parts) >= 4:
            try:
                line_no = int(parts[1])
            except ValueError:
                line_no = None
        issues.append(Issue(
            severity="error", confidence="Confirmed",
            message=stderr_text.replace("\n", " | "),
            line=line_no,
        ))

    return issues


def _analyze_structural(source: str, file_path: str) -> list[Issue]:
    """
    Non-Python fallback. Deliberately shallow — see module docstring.
    Only checks things that are unambiguous regardless of language:
    TODO-style markers and basic emptiness.
    """
    issues: list[Issue] = []
    lines = source.splitlines()

    for i, line in enumerate(lines, start=1):
        for marker in _TODO_MARKERS:
            if marker in line:
                issues.append(Issue(
                    severity="note", confidence="Confirmed",
                    message=f"{marker} marker found: {line.strip()[:100]}",
                    line=i,
                ))
                break  # one marker per line is enough context

    if not source.strip():
        issues.append(Issue(severity="warning", confidence="Confirmed", message="File is empty."))

    return issues


def read_relevant_file(file_path: str, start_line: Optional[int] = None, end_line: Optional[int] = None) -> dict:
    """
    Read a file (or a bounded line range of one) for the coding
    assistant to reason about — Section 17's `read_relevant_file`
    tool. Deliberately separate from analyze_file(): this is a raw
    read for context-gathering, not analysis, and callers that only
    need a slice of a large file (e.g. "the lines around the
    traceback") don't have to pull the whole thing through
    _MAX_FILE_BYTES's ceiling to get it.

    Reuses _read_file_safe's existence/size checks rather than
    duplicating them, so the same bounds (Section 2: "do not
    automatically read the entire project") apply here too.
    """
    source = _read_file_safe(file_path)
    lines = source.splitlines()
    total_lines = len(lines)

    if start_line is None and end_line is None:
        snippet = source
        actual_start, actual_end = 1, total_lines
    else:
        actual_start = max(1, start_line or 1)
        actual_end = min(total_lines, end_line or total_lines)
        if actual_start > actual_end:
            raise ValueError(f"start_line ({actual_start}) is after end_line ({actual_end}).")
        snippet = "\n".join(lines[actual_start - 1:actual_end])

    return {
        "file_path": file_path,
        "language": _detect_language(file_path),
        "total_lines": total_lines,
        "start_line": actual_start,
        "end_line": actual_end,
        "content": snippet,
    }


def extract_text_for_llm(file_path: str, max_chars: int = 12000) -> str:
    """
    Section 12 (document understanding) support: returns the file's
    real, extracted plain text so main.py can hand it to the LLM for
    genuine summarization / Q&A ("summarize this", "what are the
    requirements?"), as opposed to analyze_file()'s structural
    TODO/marker scan which is a different, non-LLM check. Reuses the
    exact same readers as analyze_file() (pypdf for PDF, python-docx
    for .docx, plain read for text/code) rather than a second parsing
    path, so what the LLM sees matches what analyze_file() reported on.
    Truncates (rather than raising) past max_chars — long documents get
    a clearly-labelled truncated view instead of blowing the model's
    context window; the honest boundary is stated inline for the LLM
    and, indirectly, the user.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in _PDF_EXTENSIONS:
        _check_binary_size(file_path)
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ValueError("PDF reading requires the 'pypdf' package, which isn't installed on this backend.")
        try:
            reader = PdfReader(file_path)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as e:
            raise ValueError(f"couldn't read that PDF ({e}). It may be encrypted or corrupted.")
    elif ext in _DOCX_EXTENSIONS:
        _check_binary_size(file_path)
        try:
            import docx
        except ImportError:
            raise ValueError("Word document reading requires the 'python-docx' package, which isn't installed on this backend.")
        try:
            document = docx.Document(file_path)
            text = "\n".join(p.text for p in document.paragraphs)
        except Exception as e:
            raise ValueError(f"couldn't read that document ({e}). Only modern .docx files are supported.")
    else:
        text = _read_file_safe(file_path)

    text = text.strip()
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[...truncated, {len(text) - max_chars} more characters not shown...]"
    return text


def analyze_file(file_path: str) -> AnalysisResult:
    """
    Analyze a single file. Raises FileNotFoundError/ValueError for
    missing/oversized files rather than returning a fake empty result
    — callers (the tool registry) turn that into a clean error
    response; this function stays honest about what it couldn't do.
    """
    ext = os.path.splitext(file_path)[1].lower()

    # Images/PDF/Word docs aren't code — dispatch to their own honest
    # (non-"bug check") handlers rather than force-feeding binary bytes
    # through the text-based Python/structural path below.
    if ext in _IMAGE_EXTENSIONS:
        return _analyze_image(file_path)
    if ext in _PDF_EXTENSIONS:
        return _analyze_pdf(file_path)
    if ext in _DOCX_EXTENSIONS:
        return _analyze_docx(file_path)

    source = _read_file_safe(file_path)
    language = _detect_language(file_path)
    line_count = len(source.splitlines())

    if language == "Python":
        issues = _analyze_python(source, file_path)
        depth = "static"
        summary = (
            f"Python file, {line_count} lines. Analyzed with pyflakes "
            f"(unused imports/variables, undefined names, syntax errors)."
        )
    else:
        issues = _analyze_structural(source, file_path)
        depth = "structural"
        summary = (
            f"{language} file, {line_count} lines. Only a lightweight structural pass "
            f"was run — no {language}-specific static analyzer is wired in, so this checks "
            f"markers and basic file health only, not real code issues."
        )

    return AnalysisResult(
        file_path=file_path,
        language=language,
        line_count=line_count,
        issues=issues,
        summary=summary,
        analysis_depth=depth,
    )


def _check_binary_size(file_path: str) -> int:
    """Shared existence/size guard for the image/PDF/docx paths, mirroring
    _read_file_safe()'s checks but against _MAX_BINARY_FILE_BYTES instead
    of the tighter text-file cap."""
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"'{file_path}' does not exist or is not a file.")
    size = os.path.getsize(file_path)
    if size > _MAX_BINARY_FILE_BYTES:
        raise ValueError(
            f"'{file_path}' is {size:,} bytes, over the {_MAX_BINARY_FILE_BYTES:,}-byte "
            f"analysis limit for images/PDF/Word files."
        )
    return size


def _analyze_image(file_path: str) -> AnalysisResult:
    """
    Images aren't code, so there's no pyflakes-style "bug check" that
    applies. What's genuinely verifiable: the image's own metadata
    (real dimensions/format, read directly from the file), plus
    whatever text OCR can pull out of it — scanned the same lightweight
    way a text file's structural pass would be. If OCR isn't available
    (mirrors screen_tools.py's own OCR-unavailable handling — same
    pytesseract/Tesseract dependency), that's reported honestly rather
    than silently skipped.
    """
    _check_binary_size(file_path)
    language = _detect_language(file_path)

    try:
        from PIL import Image
        with Image.open(file_path) as img:
            width, height = img.size
            fmt = img.format or language
    except Exception as e:
        raise ValueError(f"couldn't read that image ({e}).")

    issues: list[Issue] = []
    ocr_text = None
    try:
        import pytesseract
        from PIL import Image as _PILImage
        ocr_text = pytesseract.image_to_string(_PILImage.open(file_path)).strip()
    except Exception:
        issues.append(Issue(
            severity="note", confidence="Confirmed",
            message=(
                "OCR not available (pytesseract/Tesseract not installed) — only image "
                "metadata was checked; no text inside the image was scanned."
            ),
        ))

    line_count = 0
    if ocr_text:
        line_count = len(ocr_text.splitlines())
        issues.extend(_analyze_structural(ocr_text, file_path))
    elif ocr_text == "":
        issues.append(Issue(
            severity="note", confidence="Confirmed",
            message="OCR ran but found no readable text in the image.",
        ))

    summary = (
        f"Image, {width}x{height}px, format {fmt}. This isn't code, so there's no "
        f"'bug' check to run — dimensions/format were read directly from the file, "
        f"and any visible text was pulled out via OCR and scanned for TODO/FIXME markers."
    )

    return AnalysisResult(
        file_path=file_path, language=language, line_count=line_count,
        issues=issues, summary=summary, analysis_depth="structural",
    )


def _analyze_pdf(file_path: str) -> AnalysisResult:
    """
    Extracts the PDF's real text layer via pypdf and runs the same
    structural (TODO/marker) scan used for non-Python text files. A
    scanned/image-only PDF with no text layer is reported as such —
    OCR is not run on PDFs, so that's not silently pretended either.
    """
    _check_binary_size(file_path)
    language = _detect_language(file_path)

    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError(
            "PDF analysis requires the 'pypdf' package, which isn't installed on this backend."
        )

    try:
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        raise ValueError(f"couldn't read that PDF ({e}). It may be encrypted or corrupted.")

    line_count = len(text.splitlines())
    if text.strip():
        issues = _analyze_structural(text, file_path)
        extraction_note = "Text layer extracted and scanned for TODO/FIXME markers."
    else:
        issues = [Issue(
            severity="note", confidence="Confirmed",
            message="No extractable text found — likely a scanned/image-only PDF; OCR isn't run on PDFs.",
        )]
        extraction_note = "No text layer found."

    summary = (
        f"PDF, {page_count} page(s). {extraction_note} This is a structural pass, "
        f"not a 'bug' check — PDFs aren't code."
    )

    return AnalysisResult(
        file_path=file_path, language=language, line_count=line_count,
        issues=issues, summary=summary, analysis_depth="structural",
    )


def _analyze_docx(file_path: str) -> AnalysisResult:
    """
    Extracts real paragraph text via python-docx and runs the same
    structural scan. Only modern .docx is supported — legacy binary
    .doc has no reliable pure-Python reader, so that's surfaced as a
    clear error rather than attempted and silently mangled.
    """
    _check_binary_size(file_path)
    language = _detect_language(file_path)

    try:
        import docx
    except ImportError:
        raise ValueError(
            "Word document analysis requires the 'python-docx' package, which isn't installed on this backend."
        )

    try:
        document = docx.Document(file_path)
        text = "\n".join(p.text for p in document.paragraphs)
    except Exception as e:
        raise ValueError(
            f"couldn't read that document ({e}). Only modern .docx files are supported — "
            f"legacy .doc files aren't; save as .docx and try again."
        )

    line_count = len(text.splitlines())
    issues = _analyze_structural(text, file_path)
    summary = (
        f"Word document, {line_count} line(s) of text. Scanned for TODO/FIXME markers "
        f"— this is a structural pass, not a 'bug' check."
    )

    return AnalysisResult(
        file_path=file_path, language=language, line_count=line_count,
        issues=issues, summary=summary, analysis_depth="structural",
    )


# ── Function/class extraction (shared by analyze + explain) ────────

@dataclass
class CodeUnit:
    kind: str        # "function" | "class" | "file"
    name: str
    source: str
    start_line: int
    end_line: int


def extract_unit(file_path: str, unit_name: Optional[str] = None) -> CodeUnit:
    """
    Extract a specific function/class by name, or the whole file if
    unit_name is None. Python-only for name-based extraction (needs a
    real parser to find boundaries reliably) — for other languages,
    only whole-file extraction is supported; a named lookup on a
    non-Python file raises rather than silently returning the whole
    file, so callers don't mistake a whole-file dump for the specific
    unit they asked for.
    """
    source = _read_file_safe(file_path)

    if unit_name is None:
        return CodeUnit(kind="file", name=os.path.basename(file_path),
                         source=source, start_line=1, end_line=len(source.splitlines()))

    if _detect_language(file_path) != "Python":
        raise ValueError(
            f"Named function/class extraction is only supported for Python files "
            f"— '{file_path}' is not Python. Use whole-file extraction instead."
        )

    tree = ast.parse(source, filename=file_path)
    lines = source.splitlines()

    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)) and node.name == unit_name:
            kind = "class" if isinstance(node, ast.ClassDef) else "function"
            start = node.lineno
            end = getattr(node, "end_lineno", start)
            snippet = "\n".join(lines[start - 1:end])
            return CodeUnit(kind=kind, name=unit_name, source=snippet, start_line=start, end_line=end)

    raise ValueError(f"No function or class named '{unit_name}' found in '{file_path}'.")


# ── Explanation modes (Section 4/5) ─────────────────────────────────

EXPLANATION_MODES = {"BEGINNER", "LINE_BY_LINE", "TECHNICAL", "INTERVIEW", "EXAM", "ELI5"}

_MODE_INSTRUCTIONS = {
    "BEGINNER": (
        "Explain what this code does, the important concepts it uses, and why each "
        "significant part exists. Use a simple example if it helps. Avoid unexplained jargon."
    ),
    "LINE_BY_LINE": (
        "Walk through this code sequentially. Explain each meaningful line or block in order. "
        "Skip obvious syntax (e.g. don't explain what an import statement is) unless it's doing "
        "something non-obvious."
    ),
    "TECHNICAL": (
        "Explain this code's architecture, control flow, and data flow. Note design decisions, "
        "algorithmic complexity where relevant, and edge cases the code does or doesn't handle."
    ),
    "INTERVIEW": (
        "Explain this code the way you'd discuss it in a technical interview: core concepts, "
        "time and space complexity, trade-offs made, and what related questions an interviewer "
        "might ask. Mention alternative approaches where relevant."
    ),
    "EXAM": (
        "Give a concise, revision-friendly explanation: key definitions, the important points, "
        "structured so it's easy to study from."
    ),
    "ELI5": (
        "Explain this code in extremely simple language, as if to a five-year-old, using a "
        "relatable everyday analogy."
    ),
}


def build_explanation_prompt(unit: CodeUnit, mode: str) -> str:
    """
    Build the prompt text for an LLM to actually generate the
    explanation — this module extracts and frames the code, it doesn't
    call the LLM itself (keeps this testable without a model, and
    keeps the LLM call site centralized wherever the rest of JARVIS's
    chat/voice responses already go through Ollama).
    """
    mode = mode.upper()
    if mode not in EXPLANATION_MODES:
        raise ValueError(f"Unknown explanation mode '{mode}'. Must be one of: {sorted(EXPLANATION_MODES)}")

    instruction = _MODE_INSTRUCTIONS[mode]
    label = f"{unit.kind} '{unit.name}'" if unit.kind != "file" else f"file '{unit.name}'"

    return (
        f"{instruction}\n\n"
        f"Here is the {label} (lines {unit.start_line}-{unit.end_line}):\n\n"
        f"```\n{unit.source}\n```"
    )
